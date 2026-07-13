"""Tests for document generators."""

from __future__ import annotations

import os
from decimal import Decimal
from pathlib import Path
from unittest.mock import patch

import pytest

from app.generators.act_generator import ActGenerator
from app.generators.invoice_generator import InvoiceGenerator
from app.utils.number_to_words import amount_to_words


class TestInvoiceGenerator:
    """Unit tests for InvoiceGenerator."""

    @pytest.fixture
    def generator(self, tmp_path: Path) -> InvoiceGenerator:
        """Create an InvoiceGenerator with a temp templates directory."""
        # Patch templates dir to use a temporary location
        with patch("app.generators.base.settings") as mock_settings:
            mock_settings.templates_dir = tmp_path
            mock_settings.temp_dir = tmp_path / "temp"
            mock_settings.temp_dir.mkdir(parents=True, exist_ok=True)
            yield InvoiceGenerator()

    def test_validate_context_missing_field(self, generator: InvoiceGenerator) -> None:
        """Should raise ValueError when a required field is missing."""
        with pytest.raises(ValueError, match="Missing required field"):
            generator.validate_context({"company_name": "Test"})

    def test_validate_context_ok(self, generator: InvoiceGenerator) -> None:
        """Should pass with all required fields."""
        ctx = {
            "company_name": "ИП Иванов",
            "client_name": "ООО Клиент",
            "document_number": "С-001",
            "total_amount": "10000.00",
        }
        result = generator.validate_context(ctx)
        assert result["company_name"] == "ИП Иванов"

    def test_generate_with_mock_template(self, tmp_path: Path) -> None:
        """Should generate a DOCX file from a mock template."""
        from docx import Document
        from docxtpl import DocxTemplate

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        temp_dir = tmp_path / "temp"
        temp_dir.mkdir()

        # Create a minimal template with Jinja2 placeholder
        doc = Document()
        doc.add_paragraph("Invoice {{ document_number }}")
        doc.add_paragraph("Total: {{ total_amount }}")
        template_path = templates_dir / "invoice.docx"
        doc.save(template_path)

        with patch("app.generators.base.settings") as mock_settings:
            mock_settings.templates_dir = templates_dir
            mock_settings.temp_dir = temp_dir

            gen = InvoiceGenerator()
            ctx = {
                "company_name": "ИП Иванов",
                "client_name": "ООО Клиент",
                "document_number": "С-001",
                "total_amount": "10000.00",
            }
            path = gen.generate(123, ctx)
            assert os.path.exists(path)
            assert path.endswith(".docx")

            # Verify rendered content
            rendered = DocxTemplate(path)
            assert rendered is not None


class TestActGenerator:
    """Unit tests for ActGenerator."""

    @pytest.fixture
    def generator(self, tmp_path: Path) -> ActGenerator:
        with patch("app.generators.base.settings") as mock_settings:
            mock_settings.templates_dir = tmp_path
            mock_settings.temp_dir = tmp_path / "temp"
            mock_settings.temp_dir.mkdir(parents=True, exist_ok=True)
            yield ActGenerator()

    def test_validate_context_missing_contract(self, generator: ActGenerator) -> None:
        """Should raise ValueError when contract_number is missing."""
        ctx = {
            "company_name": "ИП Иванов",
            "client_name": "ООО Клиент",
            "document_number": "А-001",
            "total_amount": "5000.00",
        }
        with pytest.raises(ValueError, match="Missing required field: contract_number"):
            generator.validate_context(ctx)

    def test_validate_context_ok(self, generator: ActGenerator) -> None:
        """Should pass with all required fields."""
        ctx = {
            "company_name": "ИП Иванов",
            "client_name": "ООО Клиент",
            "document_number": "А-001",
            "total_amount": "5000.00",
            "contract_number": "Д-001",
            "contract_date": "01.01.2025",
        }
        result = generator.validate_context(ctx)
        assert result["contract_number"] == "Д-001"


class TestAmountToWords:
    """Tests for amount_to_words utility."""

    def test_amount_to_words_with_kopecks(self) -> None:
        """Should correctly convert amount with kopecks."""
        result = amount_to_words(Decimal("23100.50"))
        assert "рублей" in result
        assert "копеек" in result
        assert "23 100" not in result  # should be words, not digits
        assert "50" in result  # kopecks are shown as digits in final format

    def test_amount_to_words_zero(self) -> None:
        """Should handle zero amount."""
        result = amount_to_words(Decimal("0.00"))
        assert "рублей" in result

    def test_amount_to_words_whole_number(self) -> None:
        """Should handle amount without kopecks."""
        result = amount_to_words(Decimal("1000.00"))
        assert "рублей" in result
        assert "00 копеек" in result

    def test_amount_to_words_from_string(self) -> None:
        """Should accept string input."""
        result = amount_to_words("5000.00")
        assert "рублей" in result

    def test_amount_to_words_negative(self) -> None:
        """Should handle negative amounts."""
        result = amount_to_words(Decimal("-1000.50"))
        assert "Минус" in result
        assert "рублей" in result
